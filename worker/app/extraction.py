"""Text extraction from supported file types: TXT, PDF (text layer), DOCX.

Scanned PDFs (no extractable text) raise NeedsOCRError so the caller can mark
the document NEEDS_OCR instead of failing it.
"""
import os


class NeedsOCRError(Exception):
    """Raised when a PDF has no selectable text layer (likely scanned)."""


class UnsupportedFileError(Exception):
    """Raised for file types this MVP does not handle."""


def extract_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".txt":
        return _extract_txt(path)
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    raise UnsupportedFileError(f"unsupported extension: {ext}")


def _extract_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as fh:
        return fh.read()


def _extract_pdf(path: str) -> str:
    import fitz  # PyMuPDF

    parts = []
    with fitz.open(path) as doc:
        for page in doc:
            parts.append(page.get_text("text"))
    text = "\n".join(parts).strip()
    if not text:
        # No text layer -> almost certainly a scanned PDF. OCR is out of scope.
        raise NeedsOCRError("PDF has no extractable text layer")
    return text


def _extract_docx(path: str) -> str:
    import docx  # python-docx

    document = docx.Document(path)
    parts = [p.text for p in document.paragraphs]
    # Include table cell text, common in legal documents.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)
